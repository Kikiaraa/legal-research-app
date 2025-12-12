from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
import os
import requests
from dotenv import load_dotenv
import json
import sys
import glob
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# 加载环境变量
load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}})

# Deepseek API配置
DEEPSEEK_API_KEY = os.getenv('DEEPSEEK_API_KEY')
if not DEEPSEEK_API_KEY:
    print("警告：未设置 DEEPSEEK_API_KEY 环境变量")
    print("请在生产环境中设置此环境变量")

DEEPSEEK_API_URL = 'https://api.deepseek.com/v1/chat/completions'

# 支持的司法辖区
JURISDICTIONS = [
    "英国",
    "加拿大", 
    "法国",
    "德国",
    "西班牙",
    "爱尔兰",
    "荷兰",
    "阿根廷",
    "阿塞拜疆",
    "土耳其"
]

# 欧盟成员国（适用GDPR）
EU_COUNTRIES = ["法国", "德国", "西班牙", "爱尔兰", "荷兰"]

# 问题模板
QUESTIONS = {
    "1": {
        "title": "是否有准入要求？",
        "prompt": "根据用户选择的司法辖区，检索对应辖区的法律法规，回答是否存在准入要求。若有相关规定，请回答'有'，并给出法律依据。若没有相关规定，请回答'无'。常见的准入要求包括：1）向司法辖区数据保护机构注册登记；或2）向司法辖区数据保护机构事前通知、备案；或3）取得司法辖区数据保护机构的授权或许可；或4）向司法辖区数据保护机构缴纳费用。"
    },
    "2": {
        "title": "适用于哪些主体？",
        "prompt": "根据对应辖区的法律法规，回答哪些主体适用准入要求，并给出相应的法律依据。若有相关规定，请进一步检索规定适用于所有行业的数据控制者或数据处理者，还是特定行业的数据控制者或数据处理者。若明确规定适用于特定行业的数据控制者或数据处理者，请说明是哪些特定行业；若没有规定适用于哪些特定行业或没有明确规定，请说明适用于所有行业。"
    },
    "3": {
        "title": "是否有豁免情形？",
        "prompt": "根据对应辖区的法律法规，回答相关准入要求是否有豁免的情形。若检索到相关规定，请说明哪些情形依法被豁免，并给出法律依据。若未检索到相关规定，请说明'暂不存在豁免情形'。"
    },
    "4": {
        "title": "在哪注册登记/备案/许可/缴费申请？",
        "prompt": "根据检索的司法辖区法律法规，请回答数据控制者或数据处理者办理注册登记/备案/许可/缴费申请的地点或平台。若检索到相关规定，请优先给出申请的官方机构、官方网站、线上系统或线下窗口名称；并提供官方参考链接或联系方式，并给出相应的法律依据。若未检索到相关规定，请说明'暂无相关规定'。"
    },
    "5": {
        "title": "是否需要缴费？",
        "prompt": "根据检索的司法辖区法律法规，判断在该司法辖区内数据控制者或数据处理者是否需要向监管机构或其他机构缴纳费用（如注册费、许可费、年费等）。若有相关规定，请说明缴费的条件、金额范围、缴费周期等具体规定，并给出相应的法律依据。若无相关规定，请直接回答'暂无相关规定'。"
    },
    "6": {
        "title": "是否规定了注册登记/备案/许可/缴费证书的有效期及续展？",
        "prompt": "根据检索的司法辖区法律法规，判断并回答该司法辖区是否规定了数据控制者/数据处理者注册登记/备案/许可/缴费证书的有效期及续展。如果检索到相关规定，请概述规定的内容（包括有效期时长、续展条件、申请程序等），并给出对应条文原文或核心摘录。如果没有规定，请说明'暂无相关规定'。"
    },
    "7": {
        "title": "没有履行相应数据处理准入的法律义务，会面临什么责任？",
        "prompt": "根据检索的相关国家法律法规，判断该司法辖区法律法规中是否有规定没有履行相应数据处理准入的法律义务所面临的法律责任。如检索到相关规定，请分别说明面临的法律责任类型（包括行政处罚、刑事责任、民事责任）、行政处罚类型（包括但不限于警告、罚款、采取纠正措施、暂停/限制数据处理活动等）、刑事责任类型（包括但不限于单处或并处罚金、监禁等）、民事赔偿（包括但不限于向个人信息主体支付赔偿金等），并列明相关法律依据。如未检索到相关规定，请说明'暂无相关规定'。"
    }
}

def load_file_content(filepath):
    """加载单个文件内容，支持txt和docx格式"""
    filename = os.path.basename(filepath)
    
    try:
        if filename.lower().endswith('.txt'):
            # 读取txt文件
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                with open(filepath, 'r', encoding='gbk') as f:
                    return f.read()
        
        elif filename.lower().endswith('.docx'):
            # 读取docx文件
            doc = docx.Document(filepath)
            return '\n'.join([para.text for para in doc.paragraphs])
        
    except Exception as e:
        print(f"读取文件 {filename} 失败: {e}")
        return ""
    
    return ""

def load_knowledge_base(jurisdiction=None):
    """根据司法辖区加载对应的知识库文件"""
    import time
    start_time = time.time()
    
    knowledge_content = ""
    knowledge_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '../knowledge-base')
    
    if not os.path.exists(knowledge_dir):
        print(f"警告：知识库目录不存在: {knowledge_dir}")
        return knowledge_content
    
    if not jurisdiction or jurisdiction not in JURISDICTIONS:
        print("未指定有效的司法辖区")
        return knowledge_content
    
    print(f"开始加载 {jurisdiction} 的知识库文件...")
    
    # 1. 加载该司法辖区的所有文件（{国家}_{法规名称}.txt/docx）
    txt_pattern = os.path.join(knowledge_dir, f"{jurisdiction}_*.txt")
    docx_pattern = os.path.join(knowledge_dir, f"{jurisdiction}_*.docx")
    matching_files = glob.glob(txt_pattern) + glob.glob(docx_pattern)
    
    # 2. 如果是欧盟成员国，自动添加GDPR文件
    if jurisdiction in EU_COUNTRIES:
        gdpr_file = os.path.join(knowledge_dir, "欧盟_GDPR.docx")
        if os.path.exists(gdpr_file) and gdpr_file not in matching_files:
            matching_files.append(gdpr_file)
            print(f"  自动添加欧盟GDPR文件")
    
    if not matching_files:
        print(f"未找到 {jurisdiction} 的法律法规文件")
        return knowledge_content
    
    print(f"找到 {len(matching_files)} 个匹配文件")
    
    # 3. 加载所有匹配的文件
    for filepath in matching_files:
        filename = os.path.basename(filepath)
        file_start = time.time()
        
        content = load_file_content(filepath)
        if content:
            # 从文件名提取法规名称
            if filename.startswith(f"{jurisdiction}_"):
                law_name = filename.replace(f"{jurisdiction}_", "").rsplit('.', 1)[0]
            elif filename.startswith("欧盟_"):
                law_name = filename.replace("欧盟_", "").rsplit('.', 1)[0]
            else:
                law_name = filename.rsplit('.', 1)[0]
            
            knowledge_content += f"""

=== {filename} ===
{jurisdiction} - {law_name}

{content}
"""
            print(f"  ✓ 加载 {filename} 耗时: {time.time() - file_start:.2f}秒")
    
    elapsed_time = time.time() - start_time
    print(f"知识库加载完成，耗时: {elapsed_time:.2f}秒，内容长度: {len(knowledge_content)} 字符")
    return knowledge_content

def extract_relevant_content(knowledge_content, question_prompt, max_chars=8000):
    """
    从知识库中提取与问题相关的内容片段
    """
    # 定义关键词映射
    keyword_mapping = {
        "准入要求": ["注册", "登记", "备案", "许可", "申请", "授权", "缴费", "费用", "通知"],
        "适用主体": ["数据控制者", "数据处理者", "控制者", "处理者", "主体", "适用", "范围"],
        "豁免情形": ["豁免", "例外", "不适用", "免除", "排除"],
        "注册登记": ["注册", "登记", "备案", "申请", "机构", "平台", "网站", "系统"],
        "缴费": ["费用", "缴费", "收费", "金额", "年费", "注册费", "许可费"],
        "有效期": ["有效期", "续展", "更新", "延期", "到期", "证书"],
        "法律责任": ["责任", "处罚", "罚款", "刑事", "民事", "行政", "违法", "制裁"]
    }
    
    # 根据问题内容确定相关关键词
    relevant_keywords = []
    for category, keywords in keyword_mapping.items():
        if any(keyword in question_prompt for keyword in keywords):
            relevant_keywords.extend(keywords)
    
    # 如果没有匹配到特定关键词，使用通用关键词
    if not relevant_keywords:
        relevant_keywords = ["数据", "个人", "保护", "法", "条", "规定"]
    
    # 按文件分割知识库内容
    file_sections = knowledge_content.split('\n===')
    relevant_sections = []
    
    for section in file_sections:
        if not section.strip():
            continue
            
        # 检查是否包含相关关键词
        section_lower = section.lower()
        relevance_score = sum(1 for keyword in relevant_keywords if keyword in section_lower)
        
        if relevance_score > 0:
            # 进一步提取相关段落
            paragraphs = section.split('\n')
            relevant_paragraphs = []
            
            for para in paragraphs:
                if any(keyword in para for keyword in relevant_keywords):
                    # 包含前后文上下文
                    para_index = paragraphs.index(para)
                    start_idx = max(0, para_index - 1)
                    end_idx = min(len(paragraphs), para_index + 2)
                    context = '\n'.join(paragraphs[start_idx:end_idx])
                    relevant_paragraphs.append(context)
            
            if relevant_paragraphs:
                # 获取文件标题
                title_line = section.split('\n')[0] if section.split('\n') else "未知文件"
                section_content = f"{title_line}\n" + '\n---\n'.join(set(relevant_paragraphs))
                relevant_sections.append(section_content)
    
    # 合并相关内容并控制长度
    filtered_content = '\n\n=== '.join(relevant_sections)
    
    # 如果内容仍然过长，进行智能截断
    if len(filtered_content) > max_chars:
        # 按重要性排序（包含更多关键词的段落优先）
        sections_with_score = []
        for section in relevant_sections:
            score = sum(1 for keyword in relevant_keywords if keyword in section.lower())
            sections_with_score.append((score, section))
        
        # 按分数排序，优先保留高分内容
        sections_with_score.sort(key=lambda x: x[0], reverse=True)
        
        filtered_content = ""
        for score, section in sections_with_score:
            if len(filtered_content) + len(section) <= max_chars:
                filtered_content += f"\n\n=== {section}"
            else:
                break
    
    return filtered_content if filtered_content else knowledge_content[:max_chars]

def call_deepseek_api(prompt, knowledge_content, jurisdiction, max_retries=2):
    """调用Deepseek API - 优化版本，智能筛选相关内容，带重试机制"""
    if not DEEPSEEK_API_KEY:
        return "错误：未配置 DEEPSEEK_API_KEY 环境变量，无法调用AI服务。请联系管理员配置API密钥。"
    
    # 提取与问题相关的内容，限制为4000字符以避免请求过大
    relevant_content = extract_relevant_content(knowledge_content, prompt, max_chars=4000)
    content_length = len(relevant_content)
    
    print(f"原始内容长度: {len(knowledge_content)} 字符")
    print(f"筛选后内容长度: {content_length} 字符")
    
    headers = {
        'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
        'Content-Type': 'application/json',
        'Connection': 'close'  # 避免连接复用问题
    }
    
    system_prompt = f"""你是一个专业的法律法规检索助手。请严格根据以下{jurisdiction}的法律法规知识库内容回答问题，不要添加知识库中没有的信息。

{jurisdiction}法律法规知识库内容：
{relevant_content}

请严格围绕当前用户问题作答，仅使用直接相关的法律条文。回答格式要求：
1. 先给出明确的结论（有/无/部分适用等）
2. 简要说明具体情况
3. 最后使用'法律依据：'标签列出相关法条原文（只需要关键条款，不要全文）

注意：
- 只引用与问题直接相关的法条
- 法条过长时可适当省略，用省略号表示
- 不要使用Markdown语法
- 不要翻译法条原文
"""
    
    data = {
        'model': 'deepseek-chat',
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': prompt}
        ],
        'temperature': 0.1,
        'max_tokens': 800  # 大幅减少token数量，避免响应过大
    }
    
    # 重试机制
    for attempt in range(max_retries + 1):
        try:
            if attempt > 0:
                print(f"重试第 {attempt} 次...")
                import time
                time.sleep(2 * attempt)  # 指数退避
            
            print(f"正在调用Deepseek API (尝试 {attempt + 1}/{max_retries + 1})...")
            
            # 使用session来管理连接
            with requests.Session() as session:
                response = session.post(
                    DEEPSEEK_API_URL, 
                    headers=headers, 
                    json=data, 
                    timeout=(15, 120),  # 连接超时15秒，读取超时120秒
                    stream=False
                )
                response.raise_for_status()
            
            # 检查响应大小
            content_length_header = response.headers.get('content-length')
            if content_length_header and int(content_length_header) > 5 * 1024 * 1024:  # 5MB限制
                print(f"警告：API响应过大: {content_length_header} bytes")
                return "错误：AI服务响应数据过大，请简化问题或联系管理员。"
            
            result = response.json()
            print(f"API调用成功")
            
            # 验证响应结构
            if 'choices' not in result or not result['choices']:
                print(f"API响应格式异常: {result}")
                if attempt < max_retries:
                    continue  # 重试
                return "错误：AI服务响应格式异常，请稍后重试。"
            
            answer = result['choices'][0]['message']['content']
            print(f"获取到答案，长度: {len(answer)} 字符")
            
            # 清理大对象
            del result
            del response
            
            return answer
            
        except requests.exceptions.Timeout as e:
            print(f"API调用超时: {str(e)}")
            if attempt < max_retries:
                continue
            return "错误：AI服务响应超时，请稍后重试。"
        except requests.exceptions.RequestException as e:
            print(f"API请求失败: {str(e)}")
            if attempt < max_retries:
                continue
            return f"API请求失败: {str(e)}"
        except (KeyError, ValueError) as e:
            print(f"API响应解析错误: {str(e)}")
            if attempt < max_retries:
                continue
            return "错误：AI服务响应格式异常，请稍后重试。"
        except Exception as e:
            print(f"API调用失败: {str(e)}")
            import traceback
            traceback.print_exc()
            if attempt < max_retries:
                continue
            return f"API调用失败: {str(e)}"
    
    return "错误：多次尝试后仍然失败，请稍后重试。"

@app.route('/api/questions', methods=['GET'])
def get_questions():
    """获取所有问题列表"""
    app.logger.info('Received request for questions list')
    return jsonify(QUESTIONS)

@app.route('/api/jurisdictions', methods=['GET'])
def get_jurisdictions():
    """获取所有司法辖区列表"""
    app.logger.info('Received request for jurisdictions list')
    return jsonify(JURISDICTIONS)

@app.route('/health', methods=['GET'])
@app.route('/api/health', methods=['GET'])
def health_check():
    """健康检查端点 - 快速响应，不加载知识库"""
    return jsonify({
        'status': 'healthy',
        'service': 'legal-research-app',
        'api_configured': DEEPSEEK_API_KEY is not None
    })

@app.route('/api/debug', methods=['GET'])
def debug_info():
    """调试信息端点"""
    import sys
    knowledge_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '../knowledge-base')
    
    debug_data = {
        'python_version': sys.version,
        'api_key_configured': DEEPSEEK_API_KEY is not None,
        'api_key_length': len(DEEPSEEK_API_KEY) if DEEPSEEK_API_KEY else 0,
        'knowledge_dir_exists': os.path.exists(knowledge_dir),
        'knowledge_dir_path': knowledge_dir,
        'supported_jurisdictions': JURISDICTIONS,
        'environment': os.environ.get('RENDER', 'local'),
        'working_directory': os.getcwd()
    }
    
    # 检查知识库文件
    if os.path.exists(knowledge_dir):
        try:
            files = os.listdir(knowledge_dir)
            debug_data['knowledge_files'] = [f for f in files if f.endswith(('.txt', '.docx'))]
            debug_data['total_files'] = len(debug_data['knowledge_files'])
        except Exception as e:
            debug_data['knowledge_files_error'] = str(e)
    
    return jsonify(debug_data)

@app.route('/', methods=['GET'])
def serve_frontend():
    """提供前端页面访问"""
    return send_from_directory(os.path.join(os.path.dirname(os.path.abspath(__file__)), '../frontend'), 'index.html')

@app.route('/<path:path>', endpoint='frontend_assets')
def serve_frontend_assets(path):
    """提供前端静态文件访问"""
    return send_from_directory(os.path.join(os.path.dirname(os.path.abspath(__file__)), '../frontend'), path)

@app.route('/api/research', methods=['POST'])
def research():
    """执行法律法规检索"""
    try:
        data = request.json
        jurisdiction = data.get('jurisdiction')
        question_ids = data.get('questions', [])
        
        print(f"收到检索请求: 司法辖区={jurisdiction}, 问题={question_ids}")
        
        if not jurisdiction:
            return jsonify({'error': '请选择司法辖区'}), 400
        
        if jurisdiction not in JURISDICTIONS:
            return jsonify({'error': f'不支持的司法辖区: {jurisdiction}'}), 400
        
        if not question_ids:
            return jsonify({'error': '请选择问题'}), 400
        
        # 检查API密钥
        if not DEEPSEEK_API_KEY:
            print("错误：API密钥未配置")
            return jsonify({'error': '服务配置错误：API密钥未设置，请联系管理员'}), 500
        
        # 按问题ID的数字顺序排序，确保输出顺序正确
        question_ids = sorted(question_ids, key=lambda x: int(x))
        print(f"问题处理顺序: {question_ids}")
        
        # 加载指定司法辖区的知识库
        print(f"开始处理 {jurisdiction} 的检索请求，问题数量: {len(question_ids)}")
        knowledge_content = load_knowledge_base(jurisdiction)
        if not knowledge_content:
            print(f"错误：未找到{jurisdiction}的知识库内容")
            return jsonify({'error': f'未找到{jurisdiction}的法律法规文件，请添加以"{jurisdiction}_"开头的.txt或.docx文件'}), 404
        
        # 生成引言（简化，不调用API，直接使用固定文本）
        introduction = f"以下是基于{jurisdiction}相关法律法规的数据隐私准入制度检索结果。"
        
        results = []
        for idx, question_id in enumerate(question_ids, 1):
            if question_id in QUESTIONS:
                question = QUESTIONS[question_id]
                print(f"处理问题 {idx}/{len(question_ids)}: {question['title']}")
                
                try:
                    prompt = f"针对{jurisdiction}，{question['prompt']}。请仅回答此问题，不要涉及其他任何问题的内容。"
                    
                    # 调用AI API
                    answer = call_deepseek_api(prompt, knowledge_content, jurisdiction)
                    
                    results.append({
                        'question_id': question_id,
                        'question_title': question['title'],
                        'answer': answer
                    })
                    print(f"问题 {idx} 处理完成")
                    
                    # 强制垃圾回收，释放内存
                    import gc
                    gc.collect()
                    
                except Exception as e:
                    print(f"处理问题 {idx} 时出错: {str(e)}")
                    import traceback
                    traceback.print_exc()
                    # 即使出错也继续处理，添加错误信息
                    results.append({
                        'question_id': question_id,
                        'question_title': question['title'],
                        'answer': f"处理此问题时出现错误: {str(e)}"
                    })
                    
                    # 出错后也进行垃圾回收
                    import gc
                    gc.collect()
        
        print(f"所有问题处理完成，共 {len(results)} 个问题")
        
        # 构建报告格式
        report = f"出海目标国数据隐私准入法律检索报告\n\n具体要求请见下文\n\n(一) {jurisdiction}\n\n{introduction}"
        
        for result in results:
            report += f"\n\nQ{result['question_id']}: {result['question_title']}\nA: {result['answer']}"
        
        print(f"报告生成成功，长度: {len(report)} 字符")
        return jsonify({'report': report})
        
    except Exception as e:
        print(f"检索请求处理失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'服务器内部错误: {str(e)}'}), 500

def create_word_document(report_data, jurisdiction):
    """创建Word文档"""
    # 创建新的Word文档
    doc = docx.Document()
    
    # 设置文档标题
    title = doc.add_heading('出海目标国数据隐私准入法律检索报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加副标题
    subtitle = doc.add_paragraph(f'司法辖区：{jurisdiction}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加生成时间
    from datetime import datetime
    time_para = doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加分隔线
    doc.add_paragraph('=' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 解析报告内容
    report_lines = report_data.split('\n')
    current_section = None
    
    for line in report_lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('(一)'):
            # 主标题
            heading = doc.add_heading(line, level=1)
        elif line.startswith('Q') and ':' in line:
            # 问题标题
            question_heading = doc.add_heading(line, level=2)
        elif line.startswith('A:'):
            # 答案内容
            answer_text = line[2:].strip()  # 移除 "A:" 前缀
            
            # 分析答案结构
            if '法律依据：' in answer_text:
                # 分离主要内容和法律依据
                parts = answer_text.split('法律依据：')
                main_content = parts[0].strip()
                legal_basis = '法律依据：' + parts[1].strip() if len(parts) > 1 else ''
                
                # 添加主要内容
                if main_content:
                    content_para = doc.add_paragraph(main_content)
                    content_para.style = 'Normal'
                
                # 添加法律依据（使用不同样式）
                if legal_basis:
                    doc.add_paragraph()  # 空行
                    legal_para = doc.add_paragraph(legal_basis)
                    legal_para.style = 'Intense Quote'
            else:
                # 没有法律依据分离的情况
                content_para = doc.add_paragraph(answer_text)
                content_para.style = 'Normal'
        elif line.startswith('以下是基于'):
            # 引言
            intro_para = doc.add_paragraph(line)
            intro_para.style = 'Normal'
        elif not line.startswith('出海目标国') and not line.startswith('具体要求'):
            # 其他内容
            if line:
                doc.add_paragraph(line)
    
    # 添加页脚
    doc.add_page_break()
    footer_para = doc.add_paragraph('本报告由法律法规检索应用自动生成')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

@app.route('/api/export-word', methods=['POST'])
def export_word():
    """导出Word文档"""
    try:
        data = request.json
        report_content = data.get('report')
        jurisdiction = data.get('jurisdiction', '未知司法辖区')
        
        if not report_content:
            return jsonify({'error': '报告内容不能为空'}), 400
        
        # 创建Word文档
        doc = create_word_document(report_content, jurisdiction)
        
        # 保存到内存
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # 生成文件名
        from datetime import datetime
        filename = f"法律检索报告_{jurisdiction}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return send_file(
            doc_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"导出Word文档失败: {str(e)}")
        return jsonify({'error': f'导出失败: {str(e)}'}), 500

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5001))
    app.run(debug=False, host='0.0.0.0', port=port)